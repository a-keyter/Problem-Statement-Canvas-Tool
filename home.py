import os
import streamlit as st
from datetime import datetime, date, timedelta

import io
from docx import Document

# TEST VERSION #
local_test = False
if local_test == True:
    from apikeys import openaikey

# CLOUD VERSION #
# openaikey = st.secrets['OPENAI_API_KEY']

# AI STUFF BELOW
# ____________________

#Import LLM - Langchain
from langchain.llms import OpenAI

#Langchain Features
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain

#############################################################

################### App Layout #########################

with st.sidebar:
    st.title("AI Problem Statement Canvas")
    st.header("People don't care about solutions, they care about their problems")
    st.divider()
    st.write("This tool uses OpenAI 3.5 Turbo Completions")
    user_openai = st.text_input("Please enter your OpenAI API Key:",)

st.header('Welcome to the Problem Statement Canvas') 
st.write('This web app is a framework for rapidly identifying potential opportunities for business and commerce in our local communities and beyond.')
st.divider()

st.subheader("Let's start to think about how the world could be a better place?")

problem = st.text_input("In a few words, try to state the problem you would like to solve.")
submit = st.button("Take Action!")

if submit and problem != "":
    if local_test == False:
        openaikey = user_openai
    os.environ['OPENAI_API_KEY'] = openaikey
    llm = OpenAI(model_name="gpt-3.5-turbo-instruct", temperature=0.7)
    
    customers_template = PromptTemplate(
        input_variables= ["problem"],
        template= """

            I'm trying to build a business in the UK that solves the following problem: {problem}. \n
            
            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the ways in which people currently solve the problem, the limitations of existing solutions to the problem, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            I want to identify the 10 percent of people for whom the problem is a real pain. Give an example persona for the type of person who experiences this problem the most. Include the following details in the persona: Why does the problem affect them? How do they experience the problem? What is their age-range? What is their field of study or work? What are their interests? What is their income (GBP)? What does their regular day look like? \n

            Start the response with "Example Persona: This is (person's name). "
    """)

    customer_2_template = PromptTemplate(
        input_variables= ["problem", "generated_personas"],
        template= """

            I'm trying to build a business in the UK that solves the following problem: {problem}. \n
            
            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the ways in which people currently solve the problem, the limitations of existing solutions to the problem, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            I want to identify the 10 percent of people for whom the problem is a real pain. The personas that I have already considered include: {generated_personas}
            
            Suggest a different example persona for the type of person who experiences this problem. Include the following details in the persona: Why does the problem affect them? How do they experience the problem? What is their age-range? What is their field of study or work? What are their interests? What is their income (GBP)? What does their regular day look like? \n

            Start the response with "Example Persona #2: "
    """)

    customer_3_template = PromptTemplate(
        input_variables= ["problem", "generated_personas"],
        template= """

            I'm trying to build a business in the UK that solves the following problem: {problem}. \n
            
            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the ways in which people currently solve the problem, the limitations of existing solutions to the problem, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            I want to identify the 10 percent of people for whom the problem is a real pain. The personas that I have already considered include: {generated_personas}
            
            Suggest a different example persona for the type of person who experiences this problem. Include the following details in the persona: Why does the problem affect them? How do they experience the problem? What is their age-range? What is their field of study or work? What are their interests? What is their income (GBP)? What does their regular day look like? \n

            Start the response with "Example Persona #3: "
    """)

    context_template = PromptTemplate(
        input_variables= ['problem'],
        template= """

            I'm trying to build a business in the UK that solves the following problem: {problem}. \n
            
            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the alternatives, the opportunity, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            Give a concise response to all of the following questions about the problem:

            When does the problem typically occur for the customer? \n
            Are there specific events, actions, or conditions that trigger the problem? \n
            How does the problem escalate or become more pronounced over time? \n
            What are the indicators that suggest the problem has reached its peak in terms of pain or inconvenience for the customer? \n
            Are there any patterns or trends regarding when the customer is most likely to seek a solution or take action to address the problem? \n
    """)

    root_problem_template = PromptTemplate(
        input_variables= ['problem'],
        template= """

            I'm trying to build a business in the UK that solves the following problem: {problem}. \n
            
            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the alternatives, the opportunity, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            Suggest the reason for {problem} (First Why). Then suggest the cause for this reason under the title, Second Why. Repeat this process for a third, fourth and fifth time. 
    """)

    emotional_impact_template = PromptTemplate(
        input_variables= ['problem'],
        template= """

            I'm trying to build a business in the UK that solves the following problem: {problem}. \n
            
            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the alternatives, the opportunity, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            Suggest the emotional impact of {problem}.
    """)

    quantifiable_impact_template = PromptTemplate(
        input_variables= ['problem'],
        template= """

            I'm trying to build a business in the UK that solves the following problem: {problem}. \n
            
            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the alternatives, the opportunity, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            Suggest the quantifiable impact of {problem} at an individual level and at a societal level. Give units (metric and UK specific).
    """)

    alternatives_template = PromptTemplate(
        input_variables= ['problem', 'generated_personas'],
        template= """

            I'm trying to build a business in the UK that solves the following problem: {problem}. \n
            
            Here are 3 example personas of people that have this problem: {generated_personas}

            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the alternatives, the opportunity, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            Suggest what these people, and other people in similar circumstances, currently do to solve this problem.
    """)

    alternative_shortcomings_template = PromptTemplate(
        input_variables= ['problem', 'alternatives'],
        template= """

            I'm trying to clarify my understanding of the following problem: {problem}. \n
            
            My aim is to gain a clear understanding of the symptoms of the problem, the type of customer who would buy a potential solution, the impact of the problem, the alternatives, the opportunity, and the relationships between them, while avoiding the “solution bias” (often known as “The problem is that the customer does not use my solution". \n

            Here is an outline of the current ways that people are currently solving this problem: {alternatives}

            Suggest the limitations and shortcomings of the way in which people currently solve this problem.
    """)

    # Return the response as a json with the keys of 'when', 'triggers', 'escalation', 'peak_indicatiors', 'action_paterns'.

    #LLM Chain
    context_chain = LLMChain(llm=llm, prompt=context_template,verbose=True, output_key="context")

    customers_chain = LLMChain(llm=llm, prompt = customers_template, verbose = True, output_key="customers")
    customer_2_chain = LLMChain(llm=llm, prompt = customer_2_template, verbose=True, output_key="customer_2")
    customer_3_chain = LLMChain(llm=llm, prompt = customer_3_template, verbose=True, output_key="customer_3")

    root_problem_chain = LLMChain(llm=llm, prompt = root_problem_template, verbose = True, output_key="root_problem")

    emotional_impact_chain = LLMChain(llm=llm, prompt = emotional_impact_template, verbose = True, output_key="emotional_impact")

    quantifiable_impact_chain = LLMChain(llm=llm, prompt = quantifiable_impact_template, verbose = True, output_key="quantifiable_impact")

    alternatives_chain = LLMChain(llm=llm, prompt = alternatives_template, verbose = True, output_key="alternatives")

    alternative_shortcomings_chain = LLMChain(llm=llm, prompt = alternative_shortcomings_template, verbose = True, output_key="alternative_shortcomings")
    #_______________________________________
    #AI STUFF ABOVE

    with st.spinner("Processing"):
        context = context_chain.run({"problem": problem})
        customer_1 = customers_chain.run({"problem": problem})
        customer_2 = customer_2_chain.run({"problem": problem, "generated_personas":[customer_1]})
        customer_3 = customer_3_chain.run({"problem": problem, "generated_personas":[customer_1, customer_2]})
        root_problem = root_problem_chain.run({"problem": problem})
        emotional_impact = emotional_impact_chain.run({"problem": problem})
        quantifiable_impact = quantifiable_impact_chain.run({"problem": problem})
        alternatives = alternatives_chain.run({"problem": problem, "generated_personas": [customer_1, customer_2, customer_3]})
        alternative_shortcomings = alternative_shortcomings_chain.run({"problem": problem, "alternatives": alternatives})

    with st.expander("Customers: Who has the problem most often?"):
      st.write(customer_1)
      st.write(customer_2)
      st.write(customer_3)

    with st.expander("Context: When does the problem occur?"):
      st.write(context)

    with st.expander("Problem: What is the root cause of the problem?"):
      st.write(root_problem)

    with st.expander("Emotional Impact: How does the customer feel?"):
      st.write(emotional_impact)

    with st.expander("Quantifiable Impact: What is the measurable impact? (include units)"):
      st.write(quantifiable_impact)

    with st.expander("Alternatives: What do customers do now to fix the problem?"):
      st.write(alternatives)

    with st.expander("Alternative Shortcomings: What are the disadvantages of the alternatives?"):
      st.write(alternative_shortcomings)

    with st.spinner("Generating Report as Word Doc"):
        report = Document()

        paragraph_format = report.styles['Normal'].paragraph_format
        paragraph_format.space_after = None

        report.add_heading(f'Problem Statement Canvas: {problem}', level=1)
        
        report.add_heading("What is a Problem Statement Canvas?", level=3)
        report.add_paragraph("""
A Problem Statement Canvas is a structured tool that aids in the in-depth investigation of an issue. It’s designed to help people to clearly define a problem space and understand the problem you’re trying to solve, why it’s worth solving, whether it’s already being solved, and how solving the problem will benefit those that you seek to serve.

For more information, please see: https://www.metabeta.com/blog/process/problem-statement-canvas/
""")

        report.add_heading("What's the problem?", level=3)
        report.add_paragraph(f"This problem statement canvas aims to help the reader to understand the following problem: {problem}")

        report.add_heading("Customers: Who has the problem most often?", level=3)
        report.add_paragraph(customer_1)
        report.add_paragraph(customer_2)
        report.add_paragraph(customer_3)

        report.add_heading("Context: When does the problem occur?", level=3)
        report.add_paragraph(context)

        report.add_heading("Problem: What is the root cause of the problem?", level=3)
        report.add_paragraph(root_problem)

        report.add_heading("Emotional Impact: How does the customer feel?", level=3)
        report.add_paragraph(emotional_impact)

        report.add_heading("Quantifiable Impact: What is the measurable impact? (include units)", level=3)
        report.add_paragraph(quantifiable_impact)

        report.add_heading("Alternatives: What do customers do now to fix the problem?", level=3)
        report.add_paragraph(alternatives)

        report.add_heading("Alternative Shortcomings: What are the disadvantages of the alternatives?", level=3)
        report.add_paragraph(alternative_shortcomings)

            
        bio = io.BytesIO()
        report.save(bio)

        today_datetime = today = datetime.today()
        today_dateobj = today_datetime.date()
        today_str = str(today_dateobj)

        st.download_button(
            label="Download Canvas",
            data=bio.getvalue(),
            file_name=f"{today_str} - Problem Statement Canvas.docx",
            mime="docx",
        )

